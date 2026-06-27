"""Lesson-plans DOCX export builder.

Produces a DepEd MATATAG / ILAW lesson-plan document from a ``GeneratedLessonPlan``
and a ``SchoolHeader``. The output matches the structure of the reference template
(teacher-os-fe/LESSON PLAN.docx): centered school header paragraphs, followed by a
single ILAW table with 1 + N columns (label column + one per session).

Entry point:
    ``build_lesson_plan_docx(plan, school) -> io.BytesIO``

This function is **pure and synchronous** — it has no I/O beyond writing to an
in-memory BytesIO buffer. Callers in async contexts must offload it via
``run_in_threadpool``.
"""

import io
import re
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import _Cell

from src.lesson_plans.constants import (
    COLOR_BANNER,
    COLOR_HEADER_BG,
    COLOR_LABEL_BG,
    COLOR_TODO_BG,
    COLOR_WHITE,
    ROW_GUIDANCE,
    TEACHER_TODO_PREFIX,
)
from src.lesson_plans.schemas import GeneratedLessonPlan

FONT_NAME = "Times New Roman"


@dataclass
class SchoolHeader:
    """DepEd school information used in the top header block of the DOCX.

    All fields default to empty string so plans export cleanly even when the
    teacher has not filled in their school details yet.
    """

    school_name: str = ""
    region: str = ""
    division: str = ""
    district: str = ""
    school_address: str = ""


# ── Low-level XML helpers ─────────────────────────────────────────────────────


def _shade_cell(cell: _Cell, hex_color: str) -> None:
    """Apply a solid background fill to a table cell via the w:shd XML element.

    python-docx has no native shading API, so we inject the element directly.
    ``hex_color`` is a 6-char hex string without the leading '#'.

    Inputs: a python-docx ``_Cell`` and a 6-char hex color string.
    Outputs: none (mutates the cell's XML in place).
    Side effects: modifies ``cell._tc`` XML.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_table_borders(table) -> None:
    """Apply thin single-line borders to all table cells.

    Inputs: a python-docx ``Table``.
    Outputs: none (mutates table XML in place).
    Side effects: modifies ``table._tbl`` XML.
    """
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def _set_col_widths(table, widths_twips: list[int]) -> None:
    """Set fixed column widths (in twips) on a table.

    Inputs: a python-docx ``Table`` and a list of widths in twips (1 inch = 1440 twips).
    Outputs: none (mutates table XML in place).
    Side effects: modifies ``table._tbl`` XML.
    """
    tbl = table._tbl
    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    else:
        for col in tbl_grid.findall(qn("w:gridCol")):
            tbl_grid.remove(col)
    for w in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(w))
        tbl_grid.append(grid_col)


# ── Cell content helpers ──────────────────────────────────────────────────────


def _style_run(
    run,
    font_size: int = 12,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
) -> None:
    """Apply Times New Roman + size/bold/italic/color to a run.

    Inputs: a python-docx ``Run`` and optional formatting values.
    Outputs: none.
    Side effects: mutates run properties.
    """
    run.font.name = FONT_NAME
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _style_para(
    para,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    """Apply single (1.0) line spacing, zero space-before/after, and alignment.

    Inputs: a python-docx ``Paragraph`` and the desired alignment.
    Outputs: none.
    Side effects: mutates paragraph format properties.
    """
    para.alignment = alignment
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def _cell_para(
    cell: _Cell,
    text: str,
    bold: bool = False,
    italic: bool = False,
    font_size: int = 12,
    color: str | None = None,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """Replace a cell's first paragraph with styled text.

    Clears existing paragraphs (keeps the first for the Word container), sets
    alignment, and applies run-level formatting.

    Inputs: cell, text content, and optional run/paragraph formatting overrides.
    Outputs: none (mutates cell XML).
    Side effects: modifies cell content.
    """
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    para = cell.paragraphs[0]
    para.clear()
    _style_para(para, alignment)
    run = para.add_run(text)
    _style_run(run, font_size=font_size, bold=bold, italic=italic, color=color)


def _add_label_cell(cell: _Cell, label: str, guidance: str | None = None) -> None:
    """Fill the left label column cell: bold label + italic guidance below it.

    Inputs: the cell, the row label text, and optional ILAW guidance copy.
    Outputs: none.
    Side effects: modifies cell XML and shading.
    """
    _shade_cell(cell, COLOR_LABEL_BG)
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    para = cell.paragraphs[0]
    para.clear()
    _style_para(para, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(label)
    _style_run(run, bold=True)
    if guidance:
        g_para = cell.add_paragraph()
        _style_para(g_para, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        g_run = g_para.add_run(guidance)
        _style_run(g_run, font_size=12, italic=True)
        g_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _add_bullets(cell: _Cell, items: list[str]) -> None:
    """Write a justified bullet list into a cell (one paragraph per item, • prefix).

    Inputs: the cell and the list of string items.
    Outputs: none.
    Side effects: modifies cell content.
    """
    for i, item in enumerate(items):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.clear()
        _style_para(para)
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run(f"• {item}")
        _style_run(run)
        _maybe_shade_todo(cell, item)


def _add_numbered(cell: _Cell, items: list[str]) -> None:
    """Write a justified numbered list into a cell (1. 2. 3. …).

    Inputs: the cell and the list of string items.
    Outputs: none.
    Side effects: modifies cell content.
    """
    for i, item in enumerate(items):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.clear()
        _style_para(para)
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run(f"{i + 1}. {item}")
        _style_run(run)


def _add_text(cell: _Cell, text: str) -> None:
    """Write justified plain multiline text into a cell.

    Inputs: the cell and the text (newlines become separate paragraphs).
    Outputs: none.
    Side effects: modifies cell content.
    """
    lines = text.split("\n")
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.clear()
        _style_para(para)
        run = para.add_run(line)
        _style_run(run)


def _maybe_shade_todo(cell: _Cell, text: str) -> None:
    """Apply amber shading to a cell if its text is a teacher-to-complete placeholder.

    Inputs: the cell and the text to check.
    Outputs: none.
    Side effects: may modify cell XML shading.
    """
    if text.lstrip().startswith(TEACHER_TODO_PREFIX):
        _shade_cell(cell, COLOR_TODO_BG)


def _banner_row(
    table,
    text: str,
    col_count: int,
    bg_color: str = COLOR_BANNER,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """Add a full-width merged banner row (section headers like I — INTENTIONS).

    Inputs: the table, banner text, total column count, optional background color,
    and optional text alignment (defaults to LEFT for ILAW section bands; pass
    CENTER for the LESSON PLAN / LESSON INFORMATION banners).
    Outputs: none.
    Side effects: appends a row to the table.
    """
    row = table.add_row()
    cell = row.cells[0]
    for i in range(1, col_count):
        cell = cell.merge(row.cells[i])
    _shade_cell(cell, bg_color)
    _cell_para(cell, text, bold=True, font_size=12, color=COLOR_WHITE, alignment=alignment)


def _info_row(table, label: str, value: str, col_count: int) -> None:
    """Add a lesson-information row: label in col 0, value spanning cols 1..N.

    Inputs: the table, label, value, and total column count.
    Outputs: none.
    Side effects: appends a row to the table.
    """
    row = table.add_row()
    _add_label_cell(row.cells[0], label)
    value_cell = row.cells[1]
    if col_count > 2:
        for i in range(2, col_count):
            value_cell = value_cell.merge(row.cells[i])
    _add_text(value_cell, value)
    _maybe_shade_todo(value_cell, value)


def _full_width_row(table, label: str, guidance: str, content_fn, col_count: int) -> None:
    """Add a lesson-wide row: label+guidance in col 0, content spanning cols 1..N.

    Inputs: the table, label, guidance, a content-filling callback ``content_fn(cell)``,
    and total column count.
    Outputs: none.
    Side effects: appends a row to the table.
    """
    row = table.add_row()
    _add_label_cell(row.cells[0], label, guidance)
    value_cell = row.cells[1]
    if col_count > 2:
        for i in range(2, col_count):
            value_cell = value_cell.merge(row.cells[i])
    content_fn(value_cell)


def _per_session_row(table, label: str, guidance: str, session_cells_fn, n_sessions: int) -> None:
    """Add a per-session row: label+guidance in col 0, one cell per session.

    Inputs: the table, label, guidance, a callback ``session_cells_fn(cell, i)``
    called for each session column (0-indexed), and the session count.
    Outputs: none.
    Side effects: appends a row to the table.
    """
    row = table.add_row()
    _add_label_cell(row.cells[0], label, guidance)
    for i in range(n_sessions):
        session_cells_fn(row.cells[i + 1], i)


# ── Filename helper ───────────────────────────────────────────────────────────


def _slugify(text: str) -> str:
    """Convert a title to a safe filename stem (lowercase, hyphens, max 60 chars).

    Inputs: the title string.
    Outputs: a filesystem-safe slug.
    Side effects: none.
    """
    slug = re.sub(r"[^\w\s-]", "", text.lower())
    slug = re.sub(r"[\s_]+", "-", slug).strip("-")
    return slug[:60] or "lesson-plan"


# ── Main builder ──────────────────────────────────────────────────────────────


def build_lesson_plan_docx(plan: GeneratedLessonPlan, school: SchoolHeader) -> io.BytesIO:
    """Build an in-memory DOCX for the given ILAW lesson plan.

    Constructs a DepEd MATATAG / ILAW document matching the reference template
    (teacher-os-fe/LESSON PLAN.docx). Handles 1–5 sessions; for N > 1 a session
    header row is added so each column is labelled.

    Inputs:
        plan   — the ``GeneratedLessonPlan`` to render (may include local edits).
        school — the teacher's ``SchoolHeader`` for the top header block.
    Outputs:
        A seeked-to-zero ``io.BytesIO`` buffer containing the .docx binary.
    Side effects: none beyond building the in-memory document.
    """
    doc = Document()

    # ── Document-wide defaults (Times New Roman, single spacing, no extra space) ─
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    # ── Page margins (narrow) ─────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    # ── School header (centered paragraphs) ───────────────────────────────────
    header_lines = [
        "Department of Education",
        school.region or "",
        school.division or "",
        school.district or "",
        school.school_name or "",
        school.school_address or "",
    ]
    for i, line in enumerate(header_lines):
        if not line:
            continue
        p = doc.add_paragraph()
        _style_para(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        is_bold = i == 0 or i == 4  # "Department of Education" and school name
        run = p.add_run(line)
        _style_run(run, font_size=12, bold=is_bold)

    doc.add_paragraph()  # spacer

    # ── Build table ───────────────────────────────────────────────────────────
    n = len(plan.session_labels)
    col_count = 1 + n  # label column + one per session

    table = doc.add_table(rows=0, cols=col_count)
    _set_table_borders(table)

    # Column widths: label col ≈ 1.5", session cols share remaining width
    # Page width ≈ 8.5" - 2" margins (1" each side) = 6.5" usable = 9360 twips
    label_w = 2160  # ~1.5 inch
    remaining = 9360 - label_w
    session_w = remaining // n
    _set_col_widths(table, [label_w] + [session_w] * n)

    # ── LESSON PLAN title banner ───────────────────────────────────────────────
    _banner_row(table, "LESSON PLAN", col_count, COLOR_HEADER_BG, WD_ALIGN_PARAGRAPH.CENTER)

    # ── LESSON INFORMATION banner ─────────────────────────────────────────────
    _banner_row(table, "LESSON INFORMATION", col_count, COLOR_HEADER_BG, WD_ALIGN_PARAGRAPH.CENTER)

    # ── Lesson information rows ───────────────────────────────────────────────
    info = plan.lesson_information
    _info_row(table, "Lesson Title", info.title, col_count)
    _info_row(table, "Learning Area/s", info.learning_areas, col_count)
    _info_row(table, "Name of Teacher/s", info.teacher_name, col_count)
    _info_row(table, "Grade Level and Section", info.grade_level_and_section, col_count)
    _info_row(table, "No. of Sessions", info.sessions_label, col_count)

    # References as bullets
    ref_row = table.add_row()
    _add_label_cell(ref_row.cells[0], "References (books, websites, toolkits, etc.)")
    ref_cell = ref_row.cells[1]
    if col_count > 2:
        for i in range(2, col_count):
            ref_cell = ref_cell.merge(ref_row.cells[i])
    if info.references:
        _add_bullets(ref_cell, info.references)
    else:
        _cell_para(ref_cell, "—")

    _info_row(
        table,
        "Declaration of AI Use / See DO 003 s.2026 Annex A & DO 16 s.2026 Sec.2",
        info.ai_declaration,
        col_count,
    )

    # ── Session header row (only for multi-session plans) ─────────────────────
    if n > 1:
        hdr_row = table.add_row()
        _shade_cell(hdr_row.cells[0], COLOR_BANNER)
        _cell_para(hdr_row.cells[0], "Lesson Component", bold=True, font_size=12, color=COLOR_WHITE)
        for i, label in enumerate(plan.session_labels):
            _shade_cell(hdr_row.cells[i + 1], COLOR_BANNER)
            _cell_para(hdr_row.cells[i + 1], label, bold=True, font_size=12, color=COLOR_WHITE)

    # ── I — INTENTIONS ────────────────────────────────────────────────────────
    _banner_row(
        table,
        "I — INTENTIONS / Meaningful learning experiences are anchored in how well learners understand what they are learning and why.",
        col_count,
    )

    # Learning Competency (full-width)
    def _fill_competency(cell: _Cell) -> None:
        _add_numbered(cell, plan.learning_competency)

    _full_width_row(
        table,
        "Learning Competency",
        ROW_GUIDANCE["learning_competency"],
        _fill_competency,
        col_count,
    )

    # Learning Objectives (per session)
    def _fill_objectives(cell: _Cell, i: int) -> None:
        objs = plan.learning_objectives[i] if i < len(plan.learning_objectives) else []
        for j, obj in enumerate(objs):
            p = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
            p.clear()
            _style_para(p)
            _style_run(p.add_run("Knowledge: "), bold=True)
            _style_run(p.add_run(obj.knowledge))
            p2 = cell.add_paragraph()
            _style_para(p2)
            _style_run(p2.add_run("Skill: "), bold=True)
            _style_run(p2.add_run(obj.skill))
            p3 = cell.add_paragraph()
            _style_para(p3)
            _style_run(p3.add_run("Task: "), bold=True)
            _style_run(p3.add_run(obj.task))

    _per_session_row(
        table, "Learning Objectives", ROW_GUIDANCE["learning_objectives"], _fill_objectives, n
    )

    # Learner Context (full-width)
    def _fill_learner_context(cell: _Cell) -> None:
        _add_text(cell, plan.learner_context)
        _maybe_shade_todo(cell, plan.learner_context)

    _full_width_row(
        table, "Learner Context", ROW_GUIDANCE["learner_context"], _fill_learner_context, col_count
    )

    # ── L — LEARNING EXPERIENCES ──────────────────────────────────────────────
    _banner_row(
        table,
        "L — LEARNING EXPERIENCES / A learning experience is like a thoughtfully prepared meal — it nourishes, considers the diner’s needs, and leaves them satisfied and hungry for more.",
        col_count,
    )

    # Pre-Lesson (per session)
    def _fill_pre_lesson(cell: _Cell, i: int) -> None:
        text = plan.pre_lesson[i] if i < len(plan.pre_lesson) else ""
        _add_text(cell, text)
        _maybe_shade_todo(cell, text)

    _per_session_row(table, "Pre-Lesson", ROW_GUIDANCE["pre_lesson"], _fill_pre_lesson, n)

    # Flow (per session)
    def _fill_flow(cell: _Cell, i: int) -> None:
        steps = plan.flow[i] if i < len(plan.flow) else []
        for j, step in enumerate(steps):
            p = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
            p.clear()
            _style_para(p)
            p.paragraph_format.space_after = Pt(12)
            _style_run(p.add_run(f"{j + 1}. {step.title}: "), bold=True)
            _style_run(p.add_run(step.details))

    _per_session_row(table, "Flow / Daloy ng Aralin", ROW_GUIDANCE["flow"], _fill_flow, n)

    # Learning Resources (per session)
    def _fill_resources(cell: _Cell, i: int) -> None:
        items = plan.learning_resources[i] if i < len(plan.learning_resources) else []
        _add_bullets(cell, items)

    _per_session_row(
        table, "Learning Resources", ROW_GUIDANCE["learning_resources"], _fill_resources, n
    )

    # Opportunities for Integration (per session)
    def _fill_integration(cell: _Cell, i: int) -> None:
        items = (
            plan.opportunities_for_integration[i]
            if i < len(plan.opportunities_for_integration)
            else []
        )
        _add_bullets(cell, items)

    _per_session_row(
        table,
        "Opportunities for Integration",
        ROW_GUIDANCE["opportunities_for_integration"],
        _fill_integration,
        n,
    )

    # ── A — ASSESSMENT ────────────────────────────────────────────────────────
    _banner_row(
        table,
        "A — ASSESSMENT / Assessments reveal what learners have gained and what they still need to work on.",
        col_count,
    )

    # Formative Assessment (per session)
    def _fill_formative(cell: _Cell, i: int) -> None:
        text = plan.formative_assessment[i] if i < len(plan.formative_assessment) else ""
        _add_text(cell, text)
        _maybe_shade_todo(cell, text)

    _per_session_row(
        table, "Formative Assessment", ROW_GUIDANCE["formative_assessment"], _fill_formative, n
    )

    # ── W — WAYS FORWARD ─────────────────────────────────────────────────────
    _banner_row(
        table,
        "W — WAYS FORWARD / Meaningful learning can also happen beyond the classroom.",
        col_count,
    )

    # Extended Learning Opportunities (per session)
    def _fill_extended(cell: _Cell, i: int) -> None:
        text = (
            plan.extended_learning_opportunities[i]
            if i < len(plan.extended_learning_opportunities)
            else ""
        )
        _add_text(cell, text)
        _maybe_shade_todo(cell, text)

    _per_session_row(
        table,
        "Extended Learning Opportunities",
        ROW_GUIDANCE["extended_learning_opportunities"],
        _fill_extended,
        n,
    )

    # Reflections (per session)
    def _fill_reflections(cell: _Cell, i: int) -> None:
        text = plan.reflections[i] if i < len(plan.reflections) else ""
        _add_text(cell, text)
        _maybe_shade_todo(cell, text)

    _per_session_row(table, "Reflections", ROW_GUIDANCE["reflections"], _fill_reflections, n)

    # ── Signatories ───────────────────────────────────────────────────────────
    _banner_row(table, "PREPARED, CHECKED AND NOTED BY", col_count, COLOR_HEADER_BG)

    sig = plan.signatories
    _info_row(table, "Prepared by:", sig.prepared_by, col_count)
    _info_row(table, "Checked by:", sig.checked_by, col_count)
    _info_row(table, "Noted by:", sig.noted_by, col_count)

    # ── Serialize to buffer ───────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def docx_filename(title: str) -> str:
    """Return a safe filename for the DOCX export derived from the lesson title.

    Inputs: the lesson title string.
    Outputs: a filename string like ``introduction-to-matrices.docx``.
    Side effects: none.
    """
    return f"{_slugify(title)}.docx"
